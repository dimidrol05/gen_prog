from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox


reports_file = ""
schedule_file = ""
tekst = ""


# Функция для чтения содержимого файла
def read_docx_file():
    global reports_file
    global tekst
    root = tk.Tk()
    root.withdraw()
    reports_file = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    document = Document(reports_file)
    for paragraph in document.paragraphs:
        tekst += paragraph.text + "\n"

    return tekst


# Функция для создания таблицы расписания
def create_table(text):
    global schedule_file
    root = tk.Tk()
    root.withdraw()
    schedule_file = filedialog.asksaveasfilename(defaultextension=".docx")

    # Создаем новый документ
    document = Document()

    # Создаем таблицу с заголовками
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Время выступления"
    table.cell(0, 1).text = "ФИО"
    table.cell(0, 2).text = "Тема"

    # Разделяем текст на блоки и заполняем таблицу
    rows = text.split("\n\n")
    row_index = 1
    time = 16
    minutes = 0
    for row in rows:
        data = row.split("\n")
        if len(data) == 3 and data[2] == "Стенд":
            table.add_row().cells
            if minutes == 60:
                time += 1
                minutes = 0
            if time == 20 and minutes == 0:
                time = 16
            end_minutes = (minutes + 15) % 60
            end_time = time + 1 if end_minutes == 0 else time
            table.cell(row_index, 0).text = f"{time:02d}:{minutes:02d} - {end_time:02d}:{end_minutes:02d}"
            table.cell(row_index, 1).text = data[1]
            table.cell(row_index, 2).text = data[0]
            minutes += 15
            row_index += 1

    # Сохраняем документ и выводим сообщение о создании расписания
    document.save(schedule_file)
    messagebox.showinfo("Расписание создано", "Расписание успешно создано!")


# Функция для выбора файла с докладами
def select_reports_file():
    read_docx_file()


# Функция для выбора файла для создания расписания
def select_schedule_file():
    create_table(tekst)


# Основная функция для создания окна приложения
def main():
    root = tk.Tk()
    root.title("Приложение для создания расписания")

    # Создаем метку и кнопку для выбора файла с докладами
    select_reports_label = tk.Label(root, text="Выберите файл с докладами")
    select_reports_label.pack()
    select_reports_button = tk.Button(root, text="Выберите файл", command=select_reports_file)
    select_reports_button.pack()

    # Создаем метку и кнопку для выбора файла для создания расписания
    select_schedule_label = tk.Label(root, text="Выберите файл для создания расписания")
    select_schedule_label.pack()
    select_schedule_button = tk.Button(root, text="Выберите файл", command=select_schedule_file)
    select_schedule_button.pack()

    root.mainloop()


if __name__ == "__main__":
    main()